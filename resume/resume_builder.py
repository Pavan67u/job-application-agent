import os
import shutil
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


def read_resume_text(docx_path: str) -> str:
    """Extract plain text from DOCX resume."""
    try:
        doc = Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        return "\n".join(full_text)
    except Exception as e:
        print(f"[Resume Reader] Error: {e}")
        return ""


def save_tailored_resume(base_docx_path: str, tailored_text: str, output_dir: str, job_id: str, company: str) -> str:
    """
    Creates a new DOCX by copying base template and replacing content
    with tailored text. Returns path to new file.
    """
    os.makedirs(output_dir, exist_ok=True)
    safe_company = "".join(c for c in company if c.isalnum() or c in (" ", "_")).strip().replace(" ", "_")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    output_path = os.path.join(output_dir, f"resume_{safe_company}_{job_id}_{timestamp}.docx")

    # Copy base template to preserve formatting
    shutil.copy2(base_docx_path, output_path)
    doc = Document(output_path)

    # Clear existing paragraphs and rebuild with tailored content
    # Strategy: replace paragraph text while preserving styles
    lines = [l for l in tailored_text.split("\n") if l.strip()]
    existing_paras = [p for p in doc.paragraphs if p.text.strip()]

    # Map new content onto existing paragraphs where possible
    for i, para in enumerate(existing_paras):
        if i < len(lines):
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = lines[i]
            else:
                para.add_run(lines[i])
        else:
            # Clear extra paragraphs
            for run in para.runs:
                run.text = ""

    # Add remaining lines as new paragraphs if needed
    if len(lines) > len(existing_paras):
        for line in lines[len(existing_paras):]:
            doc.add_paragraph(line)

    doc.save(output_path)
    return output_path
